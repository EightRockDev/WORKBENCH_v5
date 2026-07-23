"""Markdown → Eight Rock-branded Word document renderer.

Built for the Artifact Engine. The LLM produces a document in the form
that's natural to it (GitHub-flavored Markdown); this module turns that
into a styled `.docx` with Eight Rock's gold/silver palette, logo cover,
DRAFT watermark, and clean typography.

Supported markdown
------------------
- Headings: `#` (Title), `##` (Section), `###` (Subsection), `####` (minor)
- Paragraphs (blank line separated, soft-wrap collapsed)
- Bold (`**text**`), italic (`*text*` or `_text_`), inline code (`` `text` ``)
- Bullet lists (`-` or `*`), numbered lists (`1.`)
- Tables (pipe-syntax with header separator row)
- Horizontal rules (`---`, `***`, `___`)
- Blockquotes (`> ...`) — italic indented
- Semantic callout blockquotes (Obsidian/GFM-style alerts):
    `> [!RECOMMENDATION]` — gold banner (used for headline calls)
    `> [!RED FLAG]` — red callout
    `> [!WATCH]` — amber callout
    `> [!INSIGHT]` — gold callout
    `> [!DD]` — blue callout (diligence item)
    `> [!NOTE]` — neutral grey callout

Inline formatting is parsed via simple regex passes over each text span;
this is "good enough for Claude's outputs" rather than a CommonMark-compliant
implementation. Edge cases not handled: nested lists deeper than 2 levels,
images, footnotes, code blocks (kept as preformatted runs).
"""

from __future__ import annotations

import datetime as dt
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import config


# ---------------------------------------------------------------------------
# Eight Rock palette (mirrors the constants in ui/exec_summary.py)
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class _Palette:
    GOLD: Any
    GOLD_DARK: Any
    GOLD_LIGHT: Any
    SILVER: Any
    CHARCOAL: Any
    BLACK: Any
    GRAY: Any
    GRAY_LIGHT: Any
    GREEN: Any
    AMBER: Any
    RED: Any
    BLUE: Any


def _palette() -> _Palette:
    from docx.shared import RGBColor
    return _Palette(
        GOLD=RGBColor(0xC8, 0x90, 0x0A),
        GOLD_DARK=RGBColor(0xA6, 0x7C, 0x00),
        GOLD_LIGHT=RGBColor(0xF7, 0xD0, 0x60),
        SILVER=RGBColor(0x80, 0x80, 0x80),
        CHARCOAL=RGBColor(0x2A, 0x2A, 0x2A),
        BLACK=RGBColor(0x11, 0x11, 0x11),
        GRAY=RGBColor(0x60, 0x60, 0x60),
        GRAY_LIGHT=RGBColor(0xC8, 0xC8, 0xC8),
        GREEN=RGBColor(0x15, 0x80, 0x3D),
        AMBER=RGBColor(0xB4, 0x53, 0x09),
        RED=RGBColor(0xB9, 0x1C, 0x1C),
        BLUE=RGBColor(0x1D, 0x4E, 0xD8),
    )


_CALLOUT_COLORS = {
    "RECOMMENDATION": ("GOLD_DARK", "Recommendation"),
    "RED FLAG": ("RED", "Red Flag"),
    "WATCH": ("AMBER", "Watch"),
    "INSIGHT": ("GOLD", "Insight"),
    "DD": ("BLUE", "Diligence Item"),
    "NOTE": ("GRAY", "Note"),
    "DD ITEM": ("BLUE", "Diligence Item"),
}


# ---------------------------------------------------------------------------
# Inline formatting (bold, italic, code) — applied to every text span
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)"
)


def _add_inline_runs(paragraph: Any, text: str, *, base_font_size: int,
                     base_color: Any) -> None:
    """Tokenize a markdown line into inline runs (plain / bold / italic /
    code) and append them to the given paragraph."""
    from docx.shared import Pt

    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        bold = italic = mono = False
        content = part
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            bold = True
            content = part[2:-2]
        elif part.startswith("__") and part.endswith("__") and len(part) >= 4:
            bold = True
            content = part[2:-2]
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            italic = True
            content = part[1:-1]
        elif part.startswith("_") and part.endswith("_") and len(part) >= 2:
            italic = True
            content = part[1:-1]
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            mono = True
            content = part[1:-1]
        run = paragraph.add_run(content)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(base_font_size)
        run.font.color.rgb = base_color
        if mono:
            run.font.name = "Consolas"


# ---------------------------------------------------------------------------
# Block-level parsing — walks the markdown line by line and emits docx
# ---------------------------------------------------------------------------

@dataclass
class _RenderContext:
    """Carries renderer state through nested helpers."""
    doc: Any
    pal: _Palette


def _add_heading(ctx: _RenderContext, level: int, text: str) -> None:
    """Add an Eight Rock-styled heading. Level maps:
    1 → 20pt gold-dark bold (used for the doc title — usually only once)
    2 → 15pt gold-dark bold (major sections)
    3 → 13pt charcoal bold (subsections)
    4 → 11pt charcoal bold-italic (minor headings inside subsections)
    """
    from docx.shared import Pt

    sizes = {1: 20, 2: 15, 3: 13, 4: 11}
    bolds = {1: True, 2: True, 3: True, 4: True}
    italics = {1: False, 2: False, 3: False, 4: True}
    colors = {
        1: ctx.pal.GOLD_DARK,
        2: ctx.pal.GOLD_DARK,
        3: ctx.pal.CHARCOAL,
        4: ctx.pal.CHARCOAL,
    }

    p = ctx.doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10 if level <= 2 else 6)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bolds.get(level, True)
    run.italic = italics.get(level, False)
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, ctx.pal.BLACK)


def _add_paragraph(ctx: _RenderContext, text: str) -> None:
    from docx.shared import Pt
    p = ctx.doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    _add_inline_runs(p, text, base_font_size=11, base_color=ctx.pal.BLACK)


def _add_bullet(ctx: _RenderContext, text: str, *, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    try:
        p = ctx.doc.add_paragraph(style=style)
    except KeyError:
        # Fallback if the style isn't registered in the default template
        p = ctx.doc.add_paragraph()
        text = ("• " if not numbered else "- ") + text
    _add_inline_runs(p, text, base_font_size=11, base_color=ctx.pal.BLACK)


def _add_blockquote(ctx: _RenderContext, lines: list[str]) -> None:
    """Generic blockquote (italic, indented, light-gray vertical bar effect
    via paragraph border)."""
    from docx.shared import Pt, Inches
    text = " ".join(line.strip() for line in lines if line.strip())
    if not text:
        return
    p = ctx.doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = ctx.pal.GRAY


def _add_callout(ctx: _RenderContext, kind: str, lines: list[str]) -> None:
    """Semantic callout block — colored left bar with label + body text."""
    from docx.shared import Pt, Inches

    color_attr, label = _CALLOUT_COLORS.get(
        kind.upper(),
        ("GRAY", kind.title()),
    )
    color = getattr(ctx.pal, color_attr, ctx.pal.GRAY)
    body_text = " ".join(line.strip() for line in lines if line.strip())
    if not body_text:
        return

    p = ctx.doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.15)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    label_run = p.add_run(f"[{label.upper()}] ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    label_run.font.color.rgb = color

    _add_inline_runs(p, body_text, base_font_size=11, base_color=ctx.pal.BLACK)


def _add_horizontal_rule(ctx: _RenderContext) -> None:
    from docx.shared import Pt
    p = ctx.doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    run = p.add_run("─" * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = ctx.pal.GRAY_LIGHT


def _add_table(ctx: _RenderContext, headers: list[str], rows: list[list[str]]) -> None:
    """Pipe-table → docx table with Eight Rock-styled header row."""
    from docx.shared import Pt
    if not headers and not rows:
        return
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return
    tbl = ctx.doc.add_table(rows=1 + len(rows), cols=n_cols)
    try:
        tbl.style = "Light Grid Accent 1"
    except KeyError:
        pass

    for i in range(n_cols):
        cell = tbl.rows[0].cells[i]
        cell.text = ""  # clear default
        text = headers[i] if i < len(headers) else ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = ctx.pal.GOLD_DARK

    for ri, row in enumerate(rows, start=1):
        for ci in range(n_cols):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            text = row[ci] if ci < len(row) else ""
            para = cell.paragraphs[0]
            _add_inline_runs(para, text, base_font_size=10, base_color=ctx.pal.BLACK)


# ---------------------------------------------------------------------------
# Top-level parser — line-driven state machine
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.+?)\s*$")
_NUMBERED_RE = re.compile(r"^(\s*)(\d+)\.\s+(.+?)\s*$")
_HR_RE = re.compile(r"^\s*([-*_])\1\1[\1\s]*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?[-:|\s]+\|?\s*$")
_CALLOUT_RE = re.compile(r"^\s*\[!\s*([A-Z][A-Z _-]*)\s*\]\s*(.*)$")


def _split_table_row(line: str) -> list[str]:
    """Pipe-row → list of cell strings. Tolerates leading/trailing pipes."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _render_markdown(doc: Any, markdown: str) -> None:
    """Walk markdown line-by-line, emitting docx elements."""
    pal = _palette()
    ctx = _RenderContext(doc=doc, pal=pal)
    lines = markdown.replace("\r\n", "\n").split("\n")

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line → paragraph break (no-op, paragraphs handle their own)
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if _HR_RE.match(line):
            _add_horizontal_rule(ctx)
            i += 1
            continue

        # Heading
        m = _HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            _add_heading(ctx, level, m.group(2).strip())
            i += 1
            continue

        # Blockquote / callout (consume consecutive `>` lines)
        if stripped.startswith(">"):
            block_lines: list[str] = []
            callout_kind: str | None = None
            callout_first_line: str | None = None
            while i < n and lines[i].strip().startswith(">"):
                content = lines[i].strip()[1:].lstrip()
                # Detect [!KIND] marker on first line
                if not block_lines and not callout_kind:
                    cm = _CALLOUT_RE.match(content)
                    if cm:
                        callout_kind = cm.group(1).strip()
                        callout_first_line = cm.group(2).strip()
                        i += 1
                        continue
                block_lines.append(content)
                i += 1
            if callout_kind is not None:
                if callout_first_line:
                    block_lines.insert(0, callout_first_line)
                _add_callout(ctx, callout_kind, block_lines)
            else:
                _add_blockquote(ctx, block_lines)
            continue

        # Table — needs at least header line + separator line
        if "|" in line and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
            headers = _split_table_row(line)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(_split_table_row(lines[i]))
                i += 1
            _add_table(ctx, headers, rows)
            continue

        # Bullet list (consume consecutive bullet lines)
        m = _BULLET_RE.match(line)
        if m:
            while i < n:
                m2 = _BULLET_RE.match(lines[i])
                if not m2:
                    break
                _add_bullet(ctx, m2.group(2))
                i += 1
            continue

        # Numbered list
        m = _NUMBERED_RE.match(line)
        if m:
            while i < n:
                m2 = _NUMBERED_RE.match(lines[i])
                if not m2:
                    break
                _add_bullet(ctx, m2.group(3), numbered=True)
                i += 1
            continue

        # Paragraph — accumulate consecutive non-blank, non-special lines
        para_lines: list[str] = []
        while i < n:
            ln = lines[i]
            if not ln.strip():
                break
            if (
                _HEADING_RE.match(ln)
                or _BULLET_RE.match(ln)
                or _NUMBERED_RE.match(ln)
                or _HR_RE.match(ln)
                or ln.strip().startswith(">")
                or ("|" in ln and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]))
            ):
                break
            para_lines.append(ln.strip())
            i += 1
        if para_lines:
            _add_paragraph(ctx, " ".join(para_lines))


# ---------------------------------------------------------------------------
# Public entry point — full doc with cover/header/watermark + body markdown
# ---------------------------------------------------------------------------

def render_markdown_to_docx(
    *,
    markdown: str,
    title: str,
    subtitle: str,
    artifact_label: str,
    property_name: str,
    property_address: str,
    output_path: Path,
    draft: bool = True,
) -> None:
    """Render a markdown body into an Eight Rock-branded Word file.

    The cover header carries: logo, optional DRAFT watermark, title,
    subtitle, property line, and a generated-at stamp. The markdown body
    flows below — full freedom for the LLM to organize as it sees fit.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    pal = _palette()
    doc = Document()

    # Tighter margins so dense memos don't waste page real estate
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Logo (PNG variant — python-docx can't render SVG)
    logo_png = (
        Path(__file__).resolve().parent.parent.parent
        / "Logos"
        / "approved-eight-rock-logo-light-preview-05062026.png"
    )
    if logo_png.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        try:
            run.add_picture(str(logo_png), width=Inches(2.4))
        except Exception:  # noqa: BLE001
            pass

    # DRAFT watermark
    if draft:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("DRAFT — NOT FOR DISTRIBUTION")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = pal.RED

    # Title
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = pal.BLACK

    # Subtitle (italic gold)
    if subtitle:
        p = doc.add_paragraph()
        r = p.add_run(subtitle)
        r.font.size = Pt(13)
        r.italic = True
        r.font.color.rgb = pal.GOLD_DARK

    # Property + address line (charcoal)
    p = doc.add_paragraph()
    r = p.add_run(property_name or "")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = pal.CHARCOAL
    if property_address:
        p = doc.add_paragraph()
        r = p.add_run(property_address)
        r.font.size = Pt(10)
        r.font.color.rgb = pal.GRAY

    # Generated stamp
    p = doc.add_paragraph()
    r = p.add_run(
        f"Generated {dt.date.today().strftime('%B %d, %Y')} · "
        f"Workbench {config.WORKBENCH_VERSION} · {artifact_label}"
    )
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = pal.GRAY

    # Thin gold rule before body
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.size = Pt(8)
    r.font.color.rgb = pal.GOLD_LIGHT

    # ---- Body markdown ----
    _render_markdown(doc, markdown)

    # Footer disclaimer
    p = doc.add_paragraph()
    r = p.add_run(
        "\nGenerated by the Eight Rock Workbench applying Beardsley/Murray/Lindahl "
        "methodology and Eight Rock's locked underwriting conventions. Every figure "
        "should be verified against source documents (T-12, rent roll, OM) before "
        "this draft is distributed."
    )
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = pal.GRAY

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
