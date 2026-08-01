"""LLM-powered analytical document generator.

Produces five different deal documents (Exec Summary, Investor Memo Summary,
Investor Memo Detail, Value-Add Strategy, LOI) — each one is GENERATED FROM
SCRATCH for the specific property, applying Brian's underwriting playbook
(Beardsley/Murray/Lindahl) and Eight Rock's locked conventions.

Architecture
------------
1. **`_build_briefing(prop, deal, folder)`** — pulls EVERYTHING the workbench
   knows about the deal into a structured dict. property record + DealState dials +
   sources.json (rent roll, T-12) + computed metrics (cap, DSCR, IRR, EM,
   cash-flow, sensitivity, refi/exit test, verdict) + market context (FMR,
   BAH, comps, lenders) + value-add levers selected. The LLM gets the WHOLE
   picture; otherwise it can't reason about cross-document consistency.

2. **`_load_guidelines()`** — reads Brian's memory files (the underwriting
   conventions, multifamily reference, value-add playbook, market intel,
   formatting rules). These become the system prompt. The LLM must apply
   them like Brian's senior analyst would.

3. **`_call_claude(...)`** — Anthropic API client. Reads `ANTHROPIC_API_KEY`
   from env (or `.env`). Returns the model's structured JSON output.

4. **`_render_to_docx(...)`** — python-docx renderer. Eight Rock branding
   (gold/silver palette, logo on cover), section headings, prose paragraphs,
   tables. Saves directly into the property folder alongside RR/T-12.

5. **`generate_artifact(...)`** — orchestrator. Returns the saved Path.

Quality controls
----------------
- Every generated doc carries a "DRAFT — NOT FOR DISTRIBUTION" header until
  Brian removes it manually. Critical for the Investor Memos where
  accidental distribution of an unreviewed draft is real risk.
- File names follow Brian's convention: `<slug>-<artifact>-MMDDYYYY.docx`.
  Regenerating on the same day overwrites; on a different day creates a
  new versioned file you can diff against.
- The LLM prompt explicitly instructs Claude to FLAG missing data rather
  than invent it (e.g., "no T-12 uploaded" not "trailing-12 NOI is roughly...").
"""

from __future__ import annotations

import datetime as dt
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import config
from core.artifact_prompts import (
    SYSTEM_PROMPT_BASE,
    prompt_for_artifact,
)
from core.calc import (
    DebtTerms,
    amortized_debt_constant,
    breakeven_occupancy,
    build_cashflow,
    build_debt_schedule,
    cap_rate,
    cash_on_cash,
    debt_yield,
    dscr,
    effective_year1_vacancy,
    return_on_cost,
)
from core.irr import project_irr
from core.risk_metrics import run_refi_exit_test
from core.sensitivity import SensitivityBase, build_sensitivity
from core.verdict import evaluate
from data.property_io import DealState, PropertyFolder, load_sources

# ---------------------------------------------------------------------------
# Artifact catalog
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ArtifactSpec:
    """One row of the artifact catalog. Defines how a given doc type is
    named, generated, and presented in the UI."""
    type_id: str          # stable identifier — used as dict key + filename slug
    label: str            # human label for the UI button
    icon: str             # emoji prefix
    description: str      # one-liner shown under the button
    enabled: bool = True  # False = stub (button visible but disabled)
    max_tokens: int = 8192


ARTIFACT_CATALOG: tuple[ArtifactSpec, ...] = (
    ArtifactSpec(
        "executive_summary",
        "Executive Summary",
        "📋",
        "Internal go/no-go memo. Analyst-blunt, threshold-cited, conditional "
        "recommendations. For Brian and Eight Rock analysts.",
        enabled=True,
        max_tokens=8192,
    ),
    ArtifactSpec(
        "investor_memo_summary",
        "Investor Memo — Trust First",
        "🤝",
        "Plain-English, personal-voice memo for the 'grandmother with $10M' "
        "investor — one who has the cash but cares whether she can trust "
        "Brian before she invests. Tax benefits explained simply, "
        "honest about what could go wrong, conservative tone. 4-6 pages.",
        enabled=True,
        max_tokens=8192,
    ),
    ArtifactSpec(
        "investor_memo_detail",
        "Investor Memo — Sophisticated",
        "📚",
        "Institutional-tone memo for the technically-fluent LP. Every "
        "mechanism walked: waterfall math, capital-call timing, tax "
        "shielding (depreciation + cost seg + 1031), exit scenarios, LP "
        "rights, governance. 15-20 pages.",
        enabled=True,
        # Bumped from 16384 → 32768 after the 16k cap truncated mid-§1031
        # before reaching the Risks (#16) and Governance (#19) sections.
        # `claude-opus-4-7` supports up to 32k output tokens natively.
        max_tokens=32768,
    ),
    ArtifactSpec(
        "value_add_strategy",
        "Value-Add Strategy",
        "🛠️",
        "Operational plan — lever-by-lever, vintage-specific, phased capex. "
        "For property mgmt + LPs. 8-12 pages.",
        enabled=True,
        max_tokens=12288,
    ),
    ArtifactSpec(
        "loi",
        "Letter of Intent",
        "✉️",
        "Legal-formal offer to seller's broker. Standard CRE LOI structure "
        "with Eight Rock terms. 2 pages.",
        enabled=True,
        max_tokens=4096,
    ),
)


def get_artifact_spec(type_id: str) -> ArtifactSpec | None:
    return next((a for a in ARTIFACT_CATALOG if a.type_id == type_id), None)


# ---------------------------------------------------------------------------
# Guidelines loader (Brian's memory files = the underwriting playbook)
# ---------------------------------------------------------------------------

# Memory files live OUTSIDE the workbench repo (in Brian's Claude Code
# memory dir). We resolve the path from $USERPROFILE so this works on any
# machine where Brian has the same Claude Code project setup.
_MEMORY_DIR = (
    Path(os.environ.get("USERPROFILE", str(Path.home())))
    / ".claude"
    / "projects"
    / (
        "C--Users-bmccu-OneDrive---Eight-Rock-Capital-Partners-"
        "8-Rock-Shared-Files-00-Technology-8-ROCK-WORKBNCH"
    )
    / "memory"
)

# Files we send to the LLM as guidelines context. Order matters — earlier
# files set the floor (locked rules) before later files add taxonomy.
_GUIDELINE_FILES = (
    "feedback_underwriting_conventions.md",     # locked GO bars, AM fee, waterfall
    "reference_multifamily_underwriting.md",    # formulas, ranges, refi/exit test
    "reference_multifamily_value_add.md",       # 21-lever playbook
    "reference_multifamily_market_intel.md",    # Hampton Roads positioning
    "feedback_financial_formatting.md",         # $X,XXX,XXX rule
    "feedback_presentation_style.md",           # voice cues
)


def _load_guidelines() -> str:
    """Concatenate all guideline memory files into a single string for the
    system prompt. Skips missing files silently — degrades gracefully if a
    file is renamed or moved."""
    blocks: list[str] = []
    for fname in _GUIDELINE_FILES:
        path = _MEMORY_DIR / fname
        if not path.is_file():
            continue
        try:
            text = path.read_text(encoding="utf-8")
        except OSError:
            continue
        blocks.append(f"# === {fname} ===\n\n{text.strip()}\n")
    return "\n\n".join(blocks)


# ---------------------------------------------------------------------------
# Briefing assembler — one canonical dict the LLM consumes
# ---------------------------------------------------------------------------

def _safe_float(v: Any, default: float | None = None) -> float | None:
    """Coerce to float, return default on any failure (None / '' / strings)."""
    if v is None or v == "":
        return default
    try:
        return float(v)
    except (TypeError, ValueError):
        return default


def _shortcut(sources: dict | None, key: str) -> dict[str, Any]:
    """Pull a `{value, source, date}` shortcut from sources.json. Returns
    empty dict if missing — the LLM uses presence to know whether the data
    is curated or inferred."""
    if not isinstance(sources, dict):
        return {}
    v = sources.get(key)
    if isinstance(v, dict) and "value" in v:
        return {
            "value": v.get("value"),
            "source": v.get("source"),
            "date": v.get("date"),
        }
    return {}


def _build_briefing(
    prop: dict[str, Any],
    deal: DealState,
    folder: PropertyFolder | None,
) -> dict[str, Any]:
    """Assemble the FULL deal briefing the LLM will reason over.

    Everything the workbench can compute or load gets surfaced here. The
    prompt explicitly instructs the model to flag missing data rather than
    invent — so passing partial briefings is OK, but every field that's
    available should be present.
    """
    sources = (
        load_sources(folder.path) if folder is not None else None
    ) or {}

    # ---- Year-1 GPR + expenses (T-12 if available, else NOI-derived) ----
    rev = sources.get("totalRevenue")
    opex = sources.get("totalOpex")
    if rev and opex:
        try:
            gpr = float(
                rev.get("value") if isinstance(rev, dict) else rev
            )
            expenses = float(
                opex.get("value") if isinstance(opex, dict) else opex
            )
        except (TypeError, ValueError):
            gpr, expenses = _gpr_from_noi(deal)
    else:
        gpr, expenses = _gpr_from_noi(deal)

    # ---- Debt schedule + 5-yr cash flow ----
    debt_terms = DebtTerms(
        loan_amount=deal.loan_amount,
        annual_rate=deal.interest_rate,
        amort_months=config.AMORT_MONTHS,
        io_years=deal.io,
    )
    debt_sched = build_debt_schedule(debt_terms, deal.hp)
    year1_eff_vac = effective_year1_vacancy(
        base_vac=deal.vacancy_frac,
        spike_pp=deal.vac_spike_pp / 100.0,
        stabilization_months=deal.stabilization_months,
    )
    cf = build_cashflow(
        year1_gpr=gpr,
        year1_vacancy_pct=year1_eff_vac,
        year1_expenses=expenses,
        rent_growth=deal.rent_growth,
        expense_growth=deal.expense_growth,
        am_fee_pct=deal.am_fee_pct,
        debt=debt_sched,
        hold_years=deal.hp,
        exit_cap=deal.exit_cap,
        equity_raise=deal.equity_raise,
        stabilized_vacancy_pct=deal.vacancy_frac,
        stabilization_year_break=1 if deal.stabilization_months <= 12 else 2,
    )

    # ---- Headline metrics ----
    cap = cap_rate(deal.noi, deal.pp)
    stab_noi = max((r.noi for r in cf.rows), default=deal.noi)
    roc = return_on_cost(stab_noi, deal.pp)
    irr_proj = project_irr(
        equity_raise=deal.equity_raise,
        annual_cashflows=[r.cash_flow for r in cf.rows],
        exit_proceeds_net=cf.exit_proceeds_net,
    )
    em = cf.equity_multiple
    ads_y1 = debt_sched.annual_payment[0]
    am_fee_y1 = cf.rows[0].am_fee if cf.rows else 0.0
    noi_after_am = deal.noi - am_fee_y1
    dscr_v = dscr(noi_after_am, ads_y1)
    coc_v = cash_on_cash(noi_after_am - ads_y1, deal.equity_raise)
    dy = debt_yield(deal.noi, deal.loan_amount)
    breakeven = (
        breakeven_occupancy(expenses, ads_y1, gpr) if gpr > 0 else 0.0
    )
    debt_const = amortized_debt_constant(deal.interest_rate, config.AMORT_MONTHS)
    ppu = (deal.pp / prop.get("units")) if prop.get("units") else None

    # ---- Refi / exit test (Beardsley) ----
    try:
        refi = run_refi_exit_test(
            stab_noi=stab_noi,
            current_balance=debt_sched.ending_balance[-1],
            current_rate=deal.interest_rate,
            cap_rate_at_exit=deal.exit_cap,
        )
    except Exception:  # noqa: BLE001 — defensive; refi math is non-critical
        refi = None

    # ---- Sensitivity grid ----
    try:
        sens = build_sensitivity(
            base=SensitivityBase(
                deal=deal, sources=sources, units=prop.get("units"),
            ),
        )
    except Exception:  # noqa: BLE001
        sens = None

    # ---- Verdict ----
    try:
        verdict = evaluate(
            cap=cap,
            dscr=dscr_v,
            coc=coc_v,
            ppu=ppu or 0.0,
            city=prop.get("city") or "",
        )
    except Exception:  # noqa: BLE001
        verdict = None

    # ---- Rent roll details ----
    rent_roll = sources.get("rentRoll") or {}
    rr_units = rent_roll.get("units") or []
    rr_summary = rent_roll.get("summary") or {}

    # ---- T-12 line items (if present) ----
    t12 = {
        k: v for k, v in sources.items()
        if k.startswith("t12_") or k in ("totalRevenue", "totalOpex")
    }

    # ---- Build the briefing ----
    return {
        "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
        "workbench_version": config.WORKBENCH_VERSION,
        "property": {
            "name": prop.get("name"),
            "address": prop.get("address"),
            "city": prop.get("city"),
            "state": prop.get("state") or "VA",
            "zip": prop.get("zip"),
            "units": prop.get("units"),
            "year_built": prop.get("year_built"),
            "last_remodel": prop.get("last_remodel"),
            "asset_class": prop.get("asset_class"),
            "property_type": prop.get("property_type"),
            "submarket": prop.get("submarket"),
            "owner": prop.get("owner"),
            "manager": prop.get("manager"),
            "management_company": prop.get("management_company"),
            "occupancy_pct_record": prop.get("occupancy_pct"),
            "avg_rent_record": prop.get("avg_rent"),
            "avg_sqft_record": prop.get("avg_sqft"),
            "rent_per_sqft_record": prop.get("rent_per_sqft"),
            "latitude": prop.get("latitude"),
            "longitude": prop.get("longitude"),
        },
        "deal_dials": {
            "purchase_price": deal.pp,
            "noi_in_place": deal.noi,
            "down_payment_pct": deal.dp,
            "interest_rate_pct": deal.ir,
            "vacancy_pct": deal.vac,
            "vacancy_source": deal.vacancy_source,
            "rent_growth_pct": deal.rg,
            "expense_growth_pct": deal.eg,
            "exit_cap_pct": deal.xc,
            "hold_period_yrs": deal.hp,
            "io_years": deal.io,
            "am_fee_pct": deal.amf,
            "equity_raise": deal.equity_raise,
            "loan_amount": deal.loan_amount,
            "tax_reassessment_on": bool(deal.tax_reassessment_on),
            "insurance_escalator_on": bool(deal.insurance_escalator_on),
            "vac_spike_pp": deal.vac_spike_pp,
            "stabilization_months": deal.stabilization_months,
            "selected_value_add_levers": list(getattr(deal, "selected_levers", []) or []),
        },
        "metrics": {
            "cap_rate_going_in": cap,
            "cap_rate_stabilized": cap_rate(stab_noi, deal.pp),
            "untrended_return_on_cost": roc,
            "dscr_year_1": dscr_v,
            "debt_yield": dy,
            "cash_on_cash_year_1": coc_v,
            "breakeven_occupancy": breakeven,
            "amortizing_debt_constant": debt_const,
            "cap_minus_debt_const_spread": cap - debt_const,
            "annual_debt_service_year_1": ads_y1,
            "year1_gpr": gpr,
            "year1_expenses": expenses,
            "year1_effective_vacancy": year1_eff_vac,
            "stabilized_noi": stab_noi,
            "exit_proceeds_net": cf.exit_proceeds_net,
            "project_irr": irr_proj,
            "equity_multiple": em,
            "price_per_unit": ppu,
        },
        "five_year_cashflow": [
            {
                "year": r.year,
                "gpr": r.gpr,
                "egi": r.egi,
                "expenses": r.expenses,
                "noi": r.noi,
                "am_fee": r.am_fee,
                "debt_service": r.debt_service,
                "cash_flow": r.cash_flow,
            }
            for r in cf.rows
        ],
        "rent_roll": {
            "source_file": rent_roll.get("file"),
            "effective_date": rent_roll.get("date"),
            "summary": rr_summary,
            "units": rr_units,
        },
        "t12": t12,
        "shortcuts": {
            "rt": _shortcut(sources, "rt"),
            "sf": _shortcut(sources, "sf"),
            "oc": _shortcut(sources, "oc"),
            "rf": _shortcut(sources, "rf"),
            "noi": _shortcut(sources, "noi"),
        },
        "refi_exit_test": _serialize(refi),
        "sensitivity": _serialize(sens),
        "verdict": _serialize(verdict),
    }


def _gpr_from_noi(deal: DealState) -> tuple[float, float]:
    """Fallback when no T-12: derive GPR + expenses from NOI + 45% Class C ER."""
    er = config.EXPENSE_RATIOS.get("C", 0.45)
    vac = deal.vacancy_frac
    denom = 1.0 - vac - er
    if denom <= 0:
        gpr = deal.noi / 0.5
    else:
        gpr = deal.noi / denom
    return gpr, gpr * er


def _serialize(obj: Any) -> Any:
    """Best-effort dataclass / pydantic / object → dict for JSON in prompt."""
    if obj is None:
        return None
    if isinstance(obj, (str, int, float, bool)):
        return obj
    if isinstance(obj, dict):
        return {k: _serialize(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_serialize(x) for x in obj]
    if hasattr(obj, "model_dump"):
        try:
            return obj.model_dump()
        except Exception:  # noqa: BLE001
            pass
    if hasattr(obj, "__dict__"):
        return {k: _serialize(v) for k, v in vars(obj).items() if not k.startswith("_")}
    return str(obj)


# ---------------------------------------------------------------------------
# Anthropic API wrapper
# ---------------------------------------------------------------------------

class ArtifactGenerationError(Exception):
    """Raised when the LLM call fails or returns malformed output."""


def _call_claude(
    *,
    system: str,
    user: str,
    model: str,
    max_tokens: int,
) -> str:
    """Single-turn Anthropic API call. Returns the response text.

    Streaming + prompt caching, with typed error classification.

    - **Streaming**: required for any request whose `max_tokens` could push
      total response time past ~10 minutes. The investor-memo-detail artifact
      uses `max_tokens=16384`, which the SDK refuses to issue non-streamed
      (Brian saw "Streaming is required for operations that may take longer
      than 10 minutes" in v0.84). `client.messages.stream(...)` solves this
      uniformly for all artifact types — no harm in streaming the smaller
      ones too.
    - **Prompt caching**: the system prompt is ~22 KB of static guidelines
      (Brian's memory files + voice rules + JSON schema), identical on every
      call. We mark it `cache_control: ephemeral` so the second-and-onward
      artifacts within the 5-min TTL serve the system prefix at ~10% of the
      input cost — a ~90% reduction on a meaningful chunk of the bill when
      Brian generates all five artifacts back-to-back for one property.
    - **Typed errors**: instead of a generic `f"Anthropic API call failed: {e}"`
      dump, we route each Anthropic exception class to an actionable message.
      The "credit balance too low" case (a 400 with a specific message
      substring) gets a one-click link to the billing page.

    Loads `.env` if present so Brian doesn't have to set the env var
    manually every shell session.
    """
    try:
        from dotenv import load_dotenv
        load_dotenv(Path(__file__).resolve().parent.parent / ".env")
    except ImportError:
        pass

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise ArtifactGenerationError(
            "ANTHROPIC_API_KEY not set. Add it to a .env file at "
            "python_workbench/.env or set the env var. See .env.example."
        )

    try:
        import anthropic
        from anthropic import Anthropic
    except ImportError as e:
        raise ArtifactGenerationError(
            "`anthropic` package not installed. Run `uv sync` (or "
            "`pip install anthropic`) and try again."
        ) from e

    # Build the system block as a list-of-text-blocks so we can attach
    # cache_control. The string form would also work but doesn't accept
    # cache_control. Default TTL is 5 minutes; bump to "ttl": "1h" if you
    # see analysts regenerating the same property an hour apart often
    # enough to justify the doubled write cost.
    system_blocks = [{
        "type": "text",
        "text": system,
        "cache_control": {"type": "ephemeral"},
    }]

    # AC-11.2: an org with ai_enabled off must reach no model at all.
    # Placed on the line that BUILDS the client, so a new surface
    # cannot forget the check and still get one.
    from core import ai_gate
    ai_gate.require_ai(
        'Artifact generation',
        'Use the deterministic Preview block, or write the document manually.',
        ai_gate.current_org_id())
    client = Anthropic(api_key=api_key)
    try:
        with client.messages.stream(
            model=model,
            max_tokens=max_tokens,
            system=system_blocks,
            messages=[{"role": "user", "content": user}],
        ) as stream:
            final_message = stream.get_final_message()
    except anthropic.BadRequestError as e:
        # 400. Three sub-cases worth distinguishing for Brian:
        #   - "credit balance too low" → billing page link
        #   - "max_tokens exceeded model limit" → likely a config bug
        #   - everything else → the raw message is usually clear enough
        msg = str(e)
        msg_lower = msg.lower()
        if "credit balance" in msg_lower or "billing" in msg_lower:
            raise ArtifactGenerationError(
                "Anthropic API credit balance is too low to generate this "
                "artifact. Add credits at "
                "https://console.anthropic.com/settings/billing then click "
                "Generate again."
            ) from e
        raise ArtifactGenerationError(
            f"Anthropic API rejected the request: {msg}"
        ) from e
    except anthropic.AuthenticationError as e:
        raise ArtifactGenerationError(
            "Anthropic API key is invalid or revoked. Check the value of "
            "ANTHROPIC_API_KEY in python_workbench/.env, or generate a new "
            "key at https://console.anthropic.com/settings/keys."
        ) from e
    except anthropic.RateLimitError as e:
        raise ArtifactGenerationError(
            "Anthropic API rate limit hit. Wait a minute and click "
            f"Generate again. ({e})"
        ) from e
    except anthropic.APIStatusError as e:
        # Generic 5xx / overloaded / etc. — the SDK already retried with
        # exponential backoff (default max_retries=2), so by the time we
        # land here the API is genuinely unhappy.
        raise ArtifactGenerationError(
            f"Anthropic API error (status {e.status_code}): {e.message}"
        ) from e
    except anthropic.APIConnectionError as e:
        raise ArtifactGenerationError(
            f"Could not reach the Anthropic API — check your internet "
            f"connection and try again. ({e})"
        ) from e

    # Walk content blocks and pick the first text block. The model can emit
    # other block types (e.g. thinking blocks if we add adaptive thinking
    # later); accessing .content[0].text without checking .type would
    # crash on those.
    if not final_message.content:
        raise ArtifactGenerationError("API returned empty response.")
    text = ""
    for block in final_message.content:
        if getattr(block, "type", None) == "text":
            text = block.text
            break
    if not text:
        raise ArtifactGenerationError(
            "API response contained no text block — got block types: "
            + ", ".join(getattr(b, "type", "?") for b in final_message.content)
        )
    return text.strip()


def _clean_markdown_response(text: str) -> str:
    """Strip any preamble before the first markdown heading or callout.

    Per the system prompt, the model MUST start with `#` or `>` — but we
    forgive minor preamble like "Here's the document:\\n\\n" defensively.
    Also strips ```markdown fences if the model wrapped its output.
    """
    s = text.strip()

    # Strip ```markdown ... ``` fences if the whole response is wrapped
    fence_match = re.match(
        r"^```(?:markdown|md)?\s*\n(.*?)\n```\s*$", s, re.DOTALL
    )
    if fence_match:
        s = fence_match.group(1).strip()

    # Find the first line starting with `#`, `>`, or text that looks like a heading
    for marker in ("\n# ", "\n## ", "\n> "):
        idx = s.find(marker)
        if idx > 0 and idx < 200:
            # There's preamble before the first markdown structure — drop it
            s = s[idx + 1 :]
            break

    # If the response starts with something other than #, ##, or >, leave as-is
    return s.strip()


# ---------------------------------------------------------------------------
# Filename + property folder helpers
# ---------------------------------------------------------------------------

def _slug(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", (name or "property").lower()).strip("-") or "property"


def _filename_for(prop: dict[str, Any], type_id: str) -> str:
    """`<property-slug>-<artifact-type>-MMDDYYYY.docx` per Brian's convention."""
    today = dt.date.today().strftime("%m%d%Y")
    parts = [
        _slug(prop.get("name") or "property"),
        type_id.replace("_", "-"),
        today,
    ]
    return "-".join(parts) + ".docx"


# ---------------------------------------------------------------------------
# DOCX renderer (legacy JSON path — kept for backwards-compat; new code
# uses `core.markdown_to_docx.render_markdown_to_docx`)
# ---------------------------------------------------------------------------

@dataclass
class _GeneratedDoc:
    """LEGACY structured output. Retained only so old tests / callers don't
    break — the live path uses markdown rendering via `markdown_to_docx`.
    """
    title: str
    subtitle: str
    sections: list[dict[str, Any]] = field(default_factory=list)
    # Optional metadata Claude can emit for footers / cover page
    headline_recommendation: str = ""
    headline_color: str = "WATCH"  # GO / WATCH / NO-GO


def _parse_generated(payload: dict[str, Any]) -> _GeneratedDoc:
    """Validate the LLM JSON shape, with forgiving defaults."""
    return _GeneratedDoc(
        title=str(payload.get("title") or "Eight Rock — Deal Memo").strip(),
        subtitle=str(payload.get("subtitle") or "").strip(),
        sections=list(payload.get("sections") or []),
        headline_recommendation=str(
            payload.get("headline_recommendation") or ""
        ).strip(),
        headline_color=str(payload.get("headline_color") or "WATCH").upper(),
    )


def _render_to_docx(
    doc_data: _GeneratedDoc,
    prop: dict[str, Any],
    artifact: ArtifactSpec,
    output_path: Path,
) -> None:
    """Render `_GeneratedDoc` to a styled Eight Rock-branded Word file."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    GOLD = RGBColor(0xC8, 0x90, 0x0A)
    GOLD_DARK = RGBColor(0xA6, 0x7C, 0x00)
    CHARCOAL = RGBColor(0x2A, 0x2A, 0x2A)
    BLACK = RGBColor(0x11, 0x11, 0x11)
    GRAY = RGBColor(0x60, 0x60, 0x60)
    GREEN = RGBColor(0x15, 0x80, 0x3D)
    AMBER = RGBColor(0xB4, 0x53, 0x09)
    RED = RGBColor(0xB9, 0x1C, 0x1C)
    headline_color = {
        "GO": GREEN, "WATCH": AMBER, "NO-GO": RED,
        "FINANCING-CONSTRAINED-WATCH": AMBER,
    }.get(doc_data.headline_color, GOLD)

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ---- Logo ----
    logo_png = (
        Path(__file__).resolve().parent.parent.parent
        / "Logos"
        / "approved-eight-rock-logo-light-preview-05062026.png"
    )
    if logo_png.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        try:
            run.add_picture(str(logo_png), width=Inches(2.4))
        except Exception:  # noqa: BLE001
            pass

    # ---- Draft watermark ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("DRAFT — NOT FOR DISTRIBUTION")
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RED

    # ---- Title ----
    p = doc.add_paragraph()
    r = p.add_run(doc_data.title)
    r.font.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLACK

    if doc_data.subtitle:
        p = doc.add_paragraph()
        r = p.add_run(doc_data.subtitle)
        r.font.size = Pt(13)
        r.font.color.rgb = GOLD_DARK
        r.font.italic = True

    # ---- Property header band ----
    name = prop.get("name") or "—"
    addr = prop.get("address") or ""
    city = prop.get("city") or ""
    state = prop.get("state") or "VA"
    zc = prop.get("zip") or ""
    p = doc.add_paragraph()
    r = p.add_run(f"{name}")
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = CHARCOAL
    p = doc.add_paragraph()
    r = p.add_run(f"{addr}, {city}, {state} {zc}".strip(", "))
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    # ---- Headline recommendation banner ----
    if doc_data.headline_recommendation:
        p = doc.add_paragraph()
        r = p.add_run(f"  {doc_data.headline_color}: ")
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = headline_color
        r = p.add_run(doc_data.headline_recommendation)
        r.font.size = Pt(11)
        r.font.color.rgb = BLACK

    # ---- Generated stamp ----
    p = doc.add_paragraph()
    r = p.add_run(
        f"Generated {dt.date.today().strftime('%B %d, %Y')} · "
        f"Workbench {config.WORKBENCH_VERSION} · "
        f"Artifact: {artifact.label}"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY
    r.font.italic = True

    # ---- Sections ----
    for section in doc_data.sections:
        heading = (section.get("heading") or "").strip()
        body = section.get("body") or []
        if heading:
            p = doc.add_paragraph()
            r = p.add_run(heading)
            r.font.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = GOLD_DARK

        for block in body if isinstance(body, list) else [body]:
            _render_block(doc, block, GOLD_DARK, BLACK, GRAY)

    # ---- Footer ----
    p = doc.add_paragraph()
    r = p.add_run(
        "\nGenerated by Eight Rock Workbench using "
        "Beardsley/Murray/Lindahl methodology and Eight Rock's locked "
        "underwriting conventions. This is a DRAFT — every figure should "
        "be verified against source documents (T-12, rent roll, OM) before "
        "distribution."
    )
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = GRAY

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _render_block(doc: Any, block: Any, accent: Any, body: Any, gray: Any) -> None:
    """Render a single content block. Supports paragraphs, bullet lists,
    sub-headings, callouts, and tables.

    Block shapes the LLM is allowed to emit:
      - {"type": "paragraph", "text": "..."}
      - {"type": "bullets", "items": ["...", "..."]}
      - {"type": "subheading", "text": "..."}
      - {"type": "callout", "label": "...", "text": "..."}
      - {"type": "table", "headers": [...], "rows": [[...], ...]}
      - bare string → treated as paragraph
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
    from docx.shared import Pt, RGBColor  # noqa: F401

    if isinstance(block, str):
        p = doc.add_paragraph()
        r = p.add_run(block)
        r.font.size = Pt(11)
        r.font.color.rgb = body
        return

    if not isinstance(block, dict):
        return

    btype = (block.get("type") or "paragraph").lower()

    if btype == "paragraph":
        p = doc.add_paragraph()
        r = p.add_run(block.get("text") or "")
        r.font.size = Pt(11)
        r.font.color.rgb = body

    elif btype == "subheading":
        p = doc.add_paragraph()
        r = p.add_run(block.get("text") or "")
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = accent

    elif btype == "bullets":
        for item in block.get("items") or []:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(item))
            r.font.size = Pt(11)
            r.font.color.rgb = body

    elif btype == "callout":
        label = (block.get("label") or "NOTE").upper()
        p = doc.add_paragraph()
        r = p.add_run(f"[{label}] ")
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = accent
        r = p.add_run(block.get("text") or "")
        r.font.size = Pt(11)
        r.font.color.rgb = body

    elif btype == "table":
        headers = block.get("headers") or []
        rows = block.get("rows") or []
        if not headers and not rows:
            return
        n_cols = max(len(headers), max((len(r) for r in rows), default=0))
        if n_cols == 0:
            return
        tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.style = "Light Grid Accent 1"
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            if i < n_cols:
                hdr_cells[i].text = str(h)
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        for ri, row_data in enumerate(rows, start=1):
            row_cells = tbl.rows[ri].cells
            for ci, val in enumerate(row_data):
                if ci < n_cols:
                    row_cells[ci].text = str(val)
                    for run in row_cells[ci].paragraphs[0].runs:
                        run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Top-level orchestrator + folder helpers
# ---------------------------------------------------------------------------

def generate_artifact(
    *,
    artifact_type: str,
    prop: dict[str, Any],
    deal: DealState,
    folder: PropertyFolder,
    # `claude-opus-4-20250514` is Opus 4.0, deprecated and retires
    # 2026-06-15. `claude-opus-4-7` is the current Opus — most capable
    # generally available model. Same input pricing ($5/M); meaningfully
    # better on long-horizon analytical work, which is exactly what these
    # artifacts are.
    model: str = "claude-opus-4-7",
) -> Path:
    """End-to-end generation. Returns the saved file path. Raises
    `ArtifactGenerationError` on any failure (caller should display).

    Pipeline (Brian's framing 2026-05-08):
      1. Python builds the perfect prompt — audience, voice, level, briefing
      2. LLM produces the document in its native form (markdown)
      3. Python parses the markdown and renders into Eight Rock Word doc
    """
    spec = get_artifact_spec(artifact_type)
    if spec is None:
        raise ArtifactGenerationError(f"Unknown artifact type: {artifact_type}")
    if not spec.enabled:
        raise ArtifactGenerationError(
            f"Artifact '{spec.label}' is not yet wired up for generation. "
            "Coming in next release."
        )
    if folder is None:
        raise ArtifactGenerationError(
            "This property has no folder yet. Upload a document or save the "
            "deal state first to create the folder."
        )

    briefing = _build_briefing(prop, deal, folder)
    guidelines = _load_guidelines()

    system_prompt = SYSTEM_PROMPT_BASE.format(
        guidelines=guidelines or "(no guideline files found)",
    )
    user_prompt = prompt_for_artifact(
        artifact_type=artifact_type,
        briefing=briefing,
    )

    raw = _call_claude(
        system=system_prompt,
        user=user_prompt,
        model=model,
        max_tokens=spec.max_tokens,
    )
    markdown_body = _clean_markdown_response(raw)
    if not markdown_body:
        raise ArtifactGenerationError(
            f"Model returned empty response. Raw:\n{raw[:300]}"
        )

    # Title / subtitle from briefing — the LLM's first heading IS the
    # document opener, but our cover banner uses property-derived text
    # for a consistent look across artifacts.
    name = prop.get("name") or "Property"
    units = prop.get("units")
    city = prop.get("city")
    cls = prop.get("asset_class")
    subtitle_bits: list[str] = [spec.label]
    if units:
        subtitle_bits.append(f"{int(units)} units")
    if cls:
        subtitle_bits.append(f"Class {cls}")
    if city:
        subtitle_bits.append(city)

    address_line = ", ".join(
        x for x in [
            prop.get("address"), prop.get("city"),
            prop.get("state") or "VA", prop.get("zip"),
        ] if x
    )

    out_path = folder.path / _filename_for(prop, artifact_type)
    from core.markdown_to_docx import render_markdown_to_docx
    render_markdown_to_docx(
        markdown=markdown_body,
        title=f"{name}",
        subtitle=" · ".join(subtitle_bits),
        artifact_label=spec.label,
        property_name=name,
        property_address=address_line,
        output_path=out_path,
        draft=True,
    )
    return out_path


def list_generated_artifacts(folder: PropertyFolder) -> dict[str, list[Path]]:
    """Scan the property folder for previously-generated artifact files,
    grouped by type_id. Used by the UI to show download buttons for past
    generations.
    """
    out: dict[str, list[Path]] = {a.type_id: [] for a in ARTIFACT_CATALOG}
    if folder is None or not folder.path.is_dir():
        return out
    for f in folder.path.iterdir():
        if not f.is_file() or f.suffix.lower() != ".docx":
            continue
        for spec in ARTIFACT_CATALOG:
            slug = spec.type_id.replace("_", "-")
            if f"-{slug}-" in f.name:
                out[spec.type_id].append(f)
                break
    for k in out:
        out[k].sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return out
