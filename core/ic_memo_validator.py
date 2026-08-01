"""IC Memo Validator — audits a generated artifact before it goes to IC.

Validates a finished `.docx` (one of the Artifact Engine's outputs) against:

  * **Structure** — required sections present per artifact type
  * **Numerics** — every canonical figure (purchase price, NOI, equity raise,
    debt service, GPR) from the briefing appears in the memo text within ±5%
  * **Verdict consistency** — the memo's stated GO/WATCH/NO-GO matches what
    `core.verdict.evaluate()` returns at the current dialed numbers
  * **DD gate** — if the memo claims IC-readiness, `due_diligence.ic_readiness()`
    must agree; if DD is gate-open, that's a flag
  * **Threshold citations** — any cited GO/WATCH/NO-GO cap bar matches current
    `core.calibration` values (within 5 bps)
  * **Voice** (optional, AI-powered) — tone alignment vs. the Brian-approved
    samples in `Templates/`. Disabled by default to avoid API spend.

Findings carry severity bands (`critical | warning | info`). A memo is
"IC-ready" when zero critical findings remain.

Consumes the same briefing dict shape as `core.artifact_engine._build_briefing()`
— pass it directly. No new database tables.
"""

from __future__ import annotations

import json
import os
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Literal

from core import calibration
from core.due_diligence import ic_readiness
from core.verdict import evaluate as evaluate_verdict

Severity = Literal["critical", "warning", "info"]


# ---------------------------------------------------------------------------
# Required-section table (per artifact type)
# ---------------------------------------------------------------------------
# Keywords are lowercase substrings checked against the lowered memo text.
# Each tuple element is one logical section; the first sub-string that matches
# counts as a hit. Multiple alternatives in a tuple are OR'd together so a
# memo can use either phrasing (e.g. "Recommendation" OR "Verdict").
_REQUIRED_SECTIONS_BY_TYPE: dict[str, tuple[tuple[str, ...], ...]] = {
    "executive_summary": (
        ("recommendation", "verdict"),
        ("purchase price", "asking price", "purchase consideration"),
        ("key metrics", "headline metrics", "summary metrics"),
        ("rationale", "thesis"),
        ("risks", "key risks", "risk factors"),
    ),
    "investor_memo_summary": (
        ("deal overview", "the deal", "property overview"),
        ("investment thesis", "why this deal"),
        ("tax benefit", "tax shielding", "depreciation"),
        ("what could go wrong", "risks", "downside"),
        ("use of funds", "sources and uses", "capital stack"),
    ),
    "investor_memo_detail": (
        ("executive summary", "deal overview"),
        ("investment thesis", "strategy"),
        ("waterfall", "distributions"),
        ("capital call", "capital plan", "funding"),
        ("tax", "depreciation", "cost segregation"),
        ("exit", "disposition"),
        ("risk", "risk factors"),
        ("governance", "lp rights", "rights"),
    ),
    "value_add_strategy": (
        ("scope of work", "scope"),
        ("phasing", "phase 1"),
        ("capex", "renovation budget"),
        ("rent premium", "premium"),
        ("timeline", "schedule"),
    ),
    "loi": (
        ("purchase price", "offer price"),
        ("deposit", "earnest money"),
        ("due diligence", "inspection period"),
        ("closing", "settlement"),
        ("exclusivity", "no-shop"),
    ),
}


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ValidationFinding:
    severity: Severity
    category: str               # "structure" | "numeric" | "verdict" | "dd_gate" | "calibration" | "voice" | "extraction"
    title: str
    message: str
    expected: Any = None
    actual: Any = None
    section: str | None = None


@dataclass
class ValidationReport:
    artifact_path: Path
    artifact_type: str
    findings: list[ValidationFinding] = field(default_factory=list)
    overall_ready: bool = False
    summary: str = ""

    @property
    def critical(self) -> list[ValidationFinding]:
        return [f for f in self.findings if f.severity == "critical"]

    @property
    def warnings(self) -> list[ValidationFinding]:
        return [f for f in self.findings if f.severity == "warning"]

    @property
    def info(self) -> list[ValidationFinding]:
        return [f for f in self.findings if f.severity == "info"]

    def to_dict(self) -> dict[str, Any]:
        return {
            "artifact_path": str(self.artifact_path),
            "artifact_type": self.artifact_type,
            "overall_ready": self.overall_ready,
            "summary": self.summary,
            "counts": {
                "critical": len(self.critical),
                "warning": len(self.warnings),
                "info": len(self.info),
            },
            "findings": [asdict(f) for f in self.findings],
        }


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_docx_text(path: Path) -> str:
    """Extract paragraph + table text from a .docx artifact. Returns "" on failure."""
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        doc = Document(str(path))
    except Exception:
        return ""
    chunks: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            chunks.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text:
                    chunks.append(cell.text)
    return "\n".join(chunks)


# ---------------------------------------------------------------------------
# Sub-validators (deterministic)
# ---------------------------------------------------------------------------

def _check_sections(text: str, artifact_type: str) -> list[ValidationFinding]:
    required = _REQUIRED_SECTIONS_BY_TYPE.get(artifact_type, ())
    lower = text.lower()
    findings: list[ValidationFinding] = []
    for alternatives in required:
        if not any(alt in lower for alt in alternatives):
            primary = alternatives[0]
            findings.append(ValidationFinding(
                severity="critical",
                category="structure",
                title=f"Missing required section: {primary.title()}",
                message=(
                    f"The {artifact_type} memo must include a '{primary}' section "
                    f"(or equivalent: {', '.join(alternatives[1:]) or 'no alternatives'}). "
                    "Per Templates voice samples."
                ),
                expected=primary,
                section=primary,
            ))
    return findings


_DOLLAR_PATTERN = re.compile(r"\$\s?(\d{1,3}(?:,\d{3})+(?:\.\d+)?)")


def _cited_dollar_figures(text: str) -> list[float]:
    """Extract all $X,XXX,XXX-style figures from memo text."""
    return [float(m.group(1).replace(",", "")) for m in _DOLLAR_PATTERN.finditer(text)]


def _check_numbers(
    text: str,
    briefing: dict[str, Any],
    tolerance: float = 0.05,
) -> list[ValidationFinding]:
    """Cross-reference canonical dollar figures from the briefing vs. memo text."""
    metrics = briefing.get("metrics") or {}
    dial = briefing.get("deal_dials") or {}
    findings: list[ValidationFinding] = []

    canonical: dict[str, float | None] = {
        "Purchase price": dial.get("purchase_price"),
        "Loan amount": dial.get("loan_amount"),
        "Equity raise": dial.get("equity_raise"),
        "NOI (in-place)": dial.get("noi_in_place"),
        "Stabilized NOI": metrics.get("stabilized_noi"),
        "Year-1 GPR": metrics.get("year1_gpr"),
        "Year-1 debt service": metrics.get("annual_debt_service_year_1"),
    }

    cited = _cited_dollar_figures(text)

    for label, value in canonical.items():
        if not isinstance(value, (int, float)) or value <= 0:
            continue
        matched = any(abs(c - value) / max(value, 1.0) <= tolerance for c in cited)
        if not matched:
            findings.append(ValidationFinding(
                severity="warning",
                category="numeric",
                title=f"{label} not found in memo (±{int(tolerance*100)}%)",
                message=(
                    f"Briefing canonical {label.lower()} is ${value:,.0f} but no figure "
                    f"within {int(tolerance*100)}% appears in the memo. Either the memo "
                    "is stale, the dial moved since generation, or the value isn't called out."
                ),
                expected=value,
            ))
    return findings


_VERDICT_EXEMPT_ARTIFACT_TYPES = frozenset({
    # The trust-first memo is written in plain English for the
    # "grandmother LP" persona — using GO/WATCH/NO-GO jargon would break
    # voice. The verdict shows up implicitly in trust-building language
    # ("a careful investment that should roughly double..."), not as a
    # token to grep for. Exempt this artifact type from verdict matching.
    "investor_memo_summary",
})


def _check_verdict_consistency(
    text: str,
    briefing: dict[str, Any],
    artifact_type: str = "",
) -> list[ValidationFinding]:
    if artifact_type in _VERDICT_EXEMPT_ARTIFACT_TYPES:
        return []
    metrics = briefing.get("metrics") or {}
    prop = briefing.get("property") or {}
    findings: list[ValidationFinding] = []

    try:
        result = evaluate_verdict(
            cap=float(metrics.get("cap_rate_going_in") or 0.0),
            dscr=float(metrics.get("dscr_year_1") or 0.0),
            coc=float(metrics.get("cash_on_cash_year_1") or 0.0),
            ppu=float(metrics.get("price_per_unit") or 0.0),
            city=prop.get("city") or "",
        )
    except Exception:
        return findings

    upper = text.upper()

    # Memo's claimed verdict — prefer the explicit `[RECOMMENDATION] <verdict>`
    # marker the prompts ask the LLM to emit. A naive substring scan over the
    # whole text falsely matches incidental mentions (e.g. "why WATCH and not
    # NO-GO: the misses are small" trips on NO-GO before the actual
    # recommendation can be read). When the marker is present, take that as
    # ground truth; when it isn't, fall back to whichever verdict token
    # dominates the memo by frequency.
    claimed: str | None = None
    marker_match = re.search(
        r"\[\s*RECOMMENDATION\s*\][^A-Z]*(FINANCING-CONSTRAINED-WATCH|NO-?GO|WATCH|GO)\b",
        upper,
    )
    if marker_match:
        token = marker_match.group(1).replace("NOGO", "NO-GO")
        claimed = token
    else:
        # Frequency-based fallback. Count each token's standalone occurrences
        # (with word boundaries so GO doesn't count inside NO-GO or NO-GO's
        # GO). Pick the most frequent; ties broken by stricter-verdict
        # priority (NO-GO > WATCH > GO).
        counts = {}
        for token, pattern in (
            ("FINANCING-CONSTRAINED-WATCH", r"FINANCING-?CONSTRAINED-?WATCH"),
            ("NO-GO", r"NO-?GO"),
            ("WATCH", r"\bWATCH\b"),
            ("GO", r"(?<!NO-)\bGO\b"),
        ):
            n = len(re.findall(pattern, upper))
            if n:
                counts[token] = n
        if counts:
            priority = {"FINANCING-CONSTRAINED-WATCH": 0, "NO-GO": 1, "WATCH": 2, "GO": 3}
            claimed = max(counts.items(), key=lambda kv: (kv[1], -priority[kv[0]]))[0]

    if claimed is None:
        findings.append(ValidationFinding(
            severity="critical",
            category="verdict",
            title="No GO/WATCH/NO-GO recommendation stated",
            message=(
                "Memo does not explicitly state a GO, WATCH, or NO-GO recommendation. "
                "Required for IC submission per Eight Rock convention."
            ),
            expected=result.verdict,
        ))
    elif claimed != result.verdict:
        findings.append(ValidationFinding(
            severity="critical",
            category="verdict",
            title="Memo verdict disagrees with calibrated verdict",
            message=(
                f"Memo claims '{claimed}' but verdict.evaluate() at current dialed numbers "
                f"returns '{result.verdict}'. Rationale: " + "; ".join(result.rationale)
            ),
            expected=result.verdict,
            actual=claimed,
        ))
    return findings


_IC_READY_PHRASES = (
    "ic-ready",
    "ic ready",
    "ready for ic",
    "ready for investment committee",
    "investment committee ready",
)


def _check_dd_gate(text: str, dd_state: Any) -> list[ValidationFinding]:
    findings: list[ValidationFinding] = []
    if dd_state is None:
        findings.append(ValidationFinding(
            severity="info",
            category="dd_gate",
            title="No DD state on file — IC-readiness gate not checked",
            message=(
                "Property has no dd.json yet. Run the Due Diligence tab to populate it; "
                "the validator will then cross-check memo's IC-ready claim against the gate."
            ),
        ))
        return findings
    try:
        readiness = ic_readiness(dd_state)
    except Exception:
        return findings

    lower = text.lower()
    claims_ready = any(p in lower for p in _IC_READY_PHRASES)

    if claims_ready and not readiness.is_ready:
        findings.append(ValidationFinding(
            severity="critical",
            category="dd_gate",
            title="Memo claims IC-ready but DD gate is open",
            message=(
                f"DD completion {readiness.completion_pct:.0%}, "
                f"{readiness.open_hard_dealbreakers} open hard dealbreakers, "
                f"{readiness.open_soft_dealbreakers_no_mitigation} soft dealbreaker(s) without "
                "mitigation. Blocking reasons: " + "; ".join(readiness.blocking_reasons)
            ),
            actual="memo says ready",
            expected=f"DD gate blocks IC (ready={readiness.is_ready})",
        ))
    elif not claims_ready and readiness.is_ready:
        findings.append(ValidationFinding(
            severity="info",
            category="dd_gate",
            title="DD gate is open — memo could explicitly state IC-readiness",
            message=(
                f"DD at {readiness.completion_pct:.0%} completion with no open hard dealbreakers. "
                "Consider adding an explicit 'IC-ready' statement to the memo."
            ),
        ))
    return findings


_THRESHOLD_PATTERN = re.compile(
    r"\b(GO|WATCH|NO-?GO)\s+(?:cap\s+(?:of|bar|floor|rate|threshold)?|bar|floor|threshold|cap)\s*[:\-]?\s*(\d+(?:\.\d+)?)\s*%",
    re.IGNORECASE,
)


def _check_threshold_citations(text: str) -> list[ValidationFinding]:
    """If the memo cites a calibrated cap threshold (e.g. 'GO cap of 7.5%'), make sure
    the cited value matches what calibration currently shows."""
    findings: list[ValidationFinding] = []
    try:
        thresholds = calibration.get_all_thresholds()
    except Exception:
        return findings
    by_name = {t.name: t for t in thresholds}

    seen: set[tuple[str, float]] = set()
    for m in _THRESHOLD_PATTERN.finditer(text):
        tier_raw = m.group(1).upper().replace("-", "")  # "NO-GO" -> "NOGO"
        cited_pct = float(m.group(2))
        if (tier_raw, cited_pct) in seen:
            continue
        seen.add((tier_raw, cited_pct))

        key = f"{tier_raw}_CAP"
        t = by_name.get(key)
        if t is None:
            continue
        calibrated_pct = t.effective_value * 100
        if abs(cited_pct - calibrated_pct) > 0.05:  # 5 bps tolerance
            tier_pretty = "NO-GO" if tier_raw == "NOGO" else tier_raw
            findings.append(ValidationFinding(
                severity="warning",
                category="calibration",
                title=f"Memo cites {tier_pretty} cap {cited_pct:.2f}% but calibration is {calibrated_pct:.2f}%",
                message=(
                    f"{key} is currently {calibrated_pct:.2f}% (source: {t.effective_source}). "
                    f"Memo states {cited_pct:.2f}% — may be stale from a prior generation, or "
                    "calibration has moved since the memo was written."
                ),
                expected=f"{calibrated_pct:.2f}%",
                actual=f"{cited_pct:.2f}%",
            ))
    return findings


# ---------------------------------------------------------------------------
# AI voice check (optional)
# ---------------------------------------------------------------------------

_VOICE_SAMPLE_FILENAMES: dict[str, str] = {
    "executive_summary": "example-executive-summary-04282026.docx",
    "investor_memo_summary": "example-investor-memo-summary-04282026.docx",
    "investor_memo_detail": "example-investor-memo-detail-04282026.docx",
    "value_add_strategy": "example-value-add-strategy-04272026.docx",
    "loi": "example-loi-05082026.docx",
}

_VOICE_SAMPLES_CACHE: dict[str, str] = {}


def _load_voice_sample(workbench_root: Path, artifact_type: str) -> str | None:
    if artifact_type in _VOICE_SAMPLES_CACHE:
        return _VOICE_SAMPLES_CACHE[artifact_type] or None
    filename = _VOICE_SAMPLE_FILENAMES.get(artifact_type)
    if not filename:
        return None
    path = workbench_root / "Templates" / filename
    if not path.is_file():
        return None
    text = extract_docx_text(path)
    _VOICE_SAMPLES_CACHE[artifact_type] = text
    return text or None


def _check_voice_with_claude(
    text: str,
    artifact_type: str,
    workbench_root: Path,
    model: str = "claude-opus-4-7",
) -> list[ValidationFinding]:
    sample = _load_voice_sample(workbench_root, artifact_type)
    if not sample:
        return [ValidationFinding(
            severity="info",
            category="voice",
            title="No voice sample available",
            message=(
                f"No Templates/example-{artifact_type.replace('_', '-')}-*.docx found. "
                "Voice check skipped."
            ),
        )]

    try:
        from dotenv import load_dotenv
        load_dotenv()
    except ImportError:
        pass
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return [ValidationFinding(
            severity="info",
            category="voice",
            title="AI voice check skipped — no ANTHROPIC_API_KEY",
            message="Set ANTHROPIC_API_KEY in .env to enable voice-alignment audit.",
        )]

    try:
        import anthropic
    except ImportError:
        return [ValidationFinding(
            severity="info",
            category="voice",
            title="AI voice check skipped — anthropic SDK not installed",
            message="Install with: uv add anthropic",
        )]

    system = (
        "You audit Eight Rock Capital IC memos for voice alignment against a Brian-approved sample. "
        "Return ONLY a JSON object with no preamble: "
        '{"issues": [{"severity": "critical"|"warning"|"info", "title": "...", "message": "..."}, ...]}.\n'
        "Flag voice/tone drift: overly promotional language, missing analyst-blunt rationale, "
        "filler corporate-speak, unsupported superlatives, structural deviations from the sample, "
        "missing risk candor, or tone inconsistent with audience for this artifact type. "
        "Do NOT flag numeric inaccuracies, structural section presence, or threshold citations — "
        "those are checked separately. If voice is well-aligned, return empty issues list."
    )
    user = (
        f"=== APPROVED SAMPLE ({artifact_type}) ===\n{sample[:8000]}\n\n"
        f"=== DRAFT TO AUDIT ===\n{text[:12000]}"
    )

    try:
        # AC-11.2: an org with ai_enabled off must reach no model at all.
        # Placed on the line that BUILDS the client, so a new surface
        # cannot forget the check and still get one.
        from core import ai_gate
        ai_gate.require_ai(
            'IC memo validation',
            'The deterministic threshold and DD-readiness checks still run.',
            ai_gate.current_org_id())
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model=model,
            max_tokens=2048,
            system=system,
            messages=[{"role": "user", "content": user}],
        )
        raw = msg.content[0].text.strip()
        # Strip code fences if Claude wrapped the JSON
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        parsed = json.loads(raw)
    except Exception as e:  # noqa: BLE001 — surface as info finding
        return [ValidationFinding(
            severity="info",
            category="voice",
            title="AI voice check failed",
            message=f"Voice audit could not run: {type(e).__name__}: {e}",
        )]

    findings: list[ValidationFinding] = []
    for issue in parsed.get("issues", []):
        sev = issue.get("severity", "info")
        if sev not in ("critical", "warning", "info"):
            sev = "info"
        findings.append(ValidationFinding(
            severity=sev,
            category="voice",
            title=str(issue.get("title", "Voice issue")),
            message=str(issue.get("message", "")),
        ))
    if not findings:
        findings.append(ValidationFinding(
            severity="info",
            category="voice",
            title="Voice aligned with approved sample",
            message="AI auditor found no tone or structural drift vs. the Templates sample.",
        ))
    return findings


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def validate(
    *,
    artifact_path: Path,
    artifact_type: str,
    briefing: dict[str, Any],
    dd_state: Any = None,
    workbench_root: Path | None = None,
    run_ai_voice_check: bool = False,
) -> ValidationReport:
    """Run the full validator suite on a generated artifact.

    Args:
        artifact_path: Path to a generated .docx memo
        artifact_type: One of the ARTIFACT_CATALOG type_ids (e.g. "executive_summary")
        briefing: The dict produced by `core.artifact_engine._build_briefing()`
        dd_state: Optional DDState; if None the DD-gate check is skipped
        workbench_root: Path to workbench root (parent of `Templates/`). Required
            only if `run_ai_voice_check=True`.
        run_ai_voice_check: When True, runs an AI voice audit against the
            Templates sample for this artifact_type. Costs API tokens.
    """
    if not artifact_path.is_file():
        return ValidationReport(
            artifact_path=artifact_path,
            artifact_type=artifact_type,
            findings=[ValidationFinding(
                severity="critical",
                category="extraction",
                title="Artifact file not found",
                message=f"No file at {artifact_path}",
            )],
            overall_ready=False,
            summary="🛑 Artifact missing — cannot validate.",
        )

    text = extract_docx_text(artifact_path)
    if not text.strip():
        return ValidationReport(
            artifact_path=artifact_path,
            artifact_type=artifact_type,
            findings=[ValidationFinding(
                severity="critical",
                category="extraction",
                title="Empty or unreadable artifact",
                message=f"Could not extract any text from {artifact_path.name}. File may be corrupt or not a valid .docx.",
            )],
            overall_ready=False,
            summary="🛑 Artifact unreadable — cannot validate.",
        )

    findings: list[ValidationFinding] = []
    findings.extend(_check_sections(text, artifact_type))
    findings.extend(_check_numbers(text, briefing))
    findings.extend(_check_verdict_consistency(text, briefing, artifact_type))
    findings.extend(_check_dd_gate(text, dd_state))
    findings.extend(_check_threshold_citations(text))
    if run_ai_voice_check and workbench_root is not None:
        findings.extend(_check_voice_with_claude(text, artifact_type, workbench_root))

    critical_count = sum(1 for f in findings if f.severity == "critical")
    warning_count = sum(1 for f in findings if f.severity == "warning")
    info_count = sum(1 for f in findings if f.severity == "info")

    overall_ready = critical_count == 0
    if overall_ready:
        summary = (
            f"✅ Ready for IC — {warning_count} warning(s), {info_count} info."
            if (warning_count or info_count)
            else "✅ Ready for IC — clean."
        )
    else:
        summary = (
            f"🛑 Revisions required — {critical_count} critical, "
            f"{warning_count} warning, {info_count} info."
        )

    return ValidationReport(
        artifact_path=artifact_path,
        artifact_type=artifact_type,
        findings=findings,
        overall_ready=overall_ready,
        summary=summary,
    )
